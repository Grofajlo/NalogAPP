import ttkbootstrap as ttkb
from ttkbootstrap.widgets import DateEntry
import sqlite3 as sq
from tkinter import BOTTOM, CENTER, NO, RIGHT, W, X, Scrollbar, Y, ttk
from tkinter import END, messagebox
from tkinter import Toplevel
import datetime
import pandas as pd
from dateutil import parser
from ttkbootstrap.constants import *
from tkinter import StringVar
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn  # for accessing XML elements in docx
from docx.oxml import OxmlElement 
from docx.shared import RGBColor
from datetime import date

#STABILNA VERZIJA APLIKACIJE ZA IZDAVANJE NALOGA

current_Y = date.today().year

con = sq.connect("Evidencija.db")
boja = RGBColor(79, 129, 189)        
headers = f'''CREATE TABLE IF NOT EXISTS Evidencija{current_Y}
                   (br_naloga INT, sektor TEXT, nalog_izdao TEXT, vođa_tima TEXT, 
                    član_1 TEXT, član_2 TEXT, član_3 TEXT, član_4 TEXT, 
                    klijent TEXT, mesto TEXT, adresa TEXT, aktivnost TEXT, datum_izdavanja DATE, rok DATE, kontakt_osoba TEXT, broj_dokumenta TEXT, invoice TEXT, napomena TEXT, rowid INT)
                  '''
con.execute(headers)

global z
z = 1

root =ttkb.Window(themename="solar")
root.title("NALOG ZA RAD")
root.grid_rowconfigure(4, weight=1)
root.grid_columnconfigure(0, weight=1)
root.iconbitmap(r"IMAGES\LOGO_TP.ico")

conn = sq.connect("Evidencija.db")
c = conn.cursor()
c.execute(f"SELECT MAX(br_naloga) FROM Evidencija{current_Y}")
maxnalog = c.fetchone()[0]

combos = []
style = ttk.Style()

# StringVar for each entry widget in client frame
firma_var = StringVar()
ponuda_var = StringVar()
kontakt_var = StringVar()
mesto_var = StringVar()
adresa_var = StringVar()
pun_naziv_var = StringVar()

con.close()

# used for making more readable reports
db_name = "Evidencija.db"

activity_map = {
    "kontrolisanje mobilnih uređaja za gašenje požara" or "Kontrolisanje mobilnih uređaja za gašenje požara": "PPA",
    "kontrolisanje instalacija hidrantske mreže za gašenje požara" or "Kontrolisanje instalacija hidrantske mreže za gašenje požara": "H",
    "kontrolisanje instalacija i uređaja za automatsko otkrivanje i dojavu požara" or "Kontrolisanje instalacija i uređaja za automatsko otkrivanje i dojavu požara": "DP",
    "Kontrolisanje PP rasvete": "PPR",
    "kontrolisanje instalacija i uređaja za gašenje požara" or "Kontrolisanje instalacija i uređaja za gašenje požara": "GP",
    "kontrolisanje instalacija i uređaja za detekciju eksplozivnih i zapaljivih gasova" or "Kontrolisanje instalacija i uređaja za detekciju eksplozivnih i zapaljivih gasova": "DETEKCIJA",
    "kontrolisanje instalacija za odvođenje dima i toplote" or "Kontrolisanje instalacija za odvođenje dima i toplote": "DIM",
    "kontrolisanje instalacija i uređaja u zonama opasnosti od eksplozije" or "Kontrolisanje instalacija i uređaja u zonama opasnosti od eksplozije": "Ex",
    "Pregled i ispitivanje električnih instalacija": "EL"
}

def shorten_activity(activity):
    for key, value in activity_map.items():
        if key in activity:
            return value
    if "mikroklime" in activity or "mikroklima" in activity:
        return "MKL"
    if "osvetljenost" in activity or "osvetljenosti" in activity:
        return "OSV"
    if "gromobranskih" in activity or "gromobrani" in activity:
        return "GR"
    if "otkrivanje" in activity or "dojavu" in activity:
        return "DP"
    if "detekciju" in activity or "eksplozivnih" in activity:
        return "Ex"
    if "hidrantske" in activity or "hidrantskih" in activity:
        return "H"
    return activity  # If no match, keep original value

# pravi izveštaj za mesec (opciono filtrira zadati mesec po zadatoj aktivnosti)
def report_v_two(month, activity_text):
    # Connect to the SQLite database
    conn = sq.connect('Evidencija.db')
    cursor = conn.cursor()
    
    # SQL query to filter data
    if activity_text:
        table_name = get_table_name()
        query = f"""
        SELECT * FROM {table_name}
        WHERE substr(rok, 4, 2) = ?
        AND aktivnost LIKE ?
        """
        cursor.execute(query, (f'{int(month):02}', f'%{activity_text}%'))
    else:
        table_name = get_table_name()        
        query = f"""
        SELECT * FROM {table_name}
        WHERE substr(rok, 4, 2) = ?
        """
        cursor.execute(query, (f'{int(month):02}',))
    
    # Fetch all filtered data
    rows = cursor.fetchall()
    
    # Get column names
    col_names = [description[0] for description in cursor.description]
    
    conn.close()
    
    # Convert data to a pandas DataFrame
    df = pd.DataFrame(rows, columns=col_names)
    
    # path for saving monthly reports
    folder_path = r"MONTHLY"
    # Write the DataFrame to an Excel file
    excel_filename = f'{folder_path}\\Izveštaj za {month} mesec.xlsx'
    df.to_excel(excel_filename, index=False)
    messagebox.showinfo(title="OBAVEŠTENJE", message=f"Izveštaj za {month}. mesec" + " je uspešno formiran")

def on_button_click():
    table_name = get_table_name()
    selected_month = mesec_combo.get()
    activity_text = search_activity_Entry.get()
    report_v_two(selected_month, activity_text)

def menu_5(): #ovo će biti za dugme x (menjaće lokacije snimanja izveštaja i naloga)
    global lvl2_5root, nalog_opseg
    lvl2_5root = ttkb.Toplevel()
    lvl2_5root.title("ŠTAMPANJE NALOGA")
    frame5 = ttkb.LabelFrame(lvl2_5root, bootstyle=SUCCESS, text="NALOG U WORD ")
    frame5.grid(row=0, column=0,padx=10, pady=10, sticky="news")
    frame4_1 = ttkb.LabelFrame(lvl2_5root, bootstyle=SUCCESS, text="NAPOMENA: ")
    frame4_1.grid(row=0,column=1, rowspan=2, padx=10, pady=10, sticky="news")
    label1 = ttkb.Label(frame5, text="BROJ NALOGA: \nILI OPSEG: ")
    label1.grid(row=0, column=0, padx=5, pady=10, sticky=E)
    label2 = ttkb.Label(frame4_1, text="U polje 'BROJ NALOGA' možete upisati broj \npojedinačnog naloga (e.g. 1982), \nili opseg koji uključuje \nobe granične vrednosti (e.g. 1982-1985), \nna ovaj način generisaće se word za \n1982, 1983, 1984 i 1985")
    label2.grid(row=0, column=2, rowspan=2)
    nalog_opseg = ttkb.Entry(frame5, width=15)
    nalog_opseg.grid(row=0, column=1, padx=5, pady=10)

    submit_button = ttkb.Button(frame5, text="BROWSE", bootstyle=SUCCESS, width=15)
    submit_button.grid(row=1, column=0, padx=5, pady=15)
    exit4_btn = ttkb.Button(frame5, text="ZATVORI I SAČUVAJ", bootstyle=SUCCESS, width=15)
    exit4_btn.grid(row=1, column=1, padx=5, pady=15)

def menu_4():
    global lvl2_4root, nalog_opseg
   
    lvl2_4root = ttkb.Toplevel()
    lvl2_4root.title("ŠTAMPANJE NALOGA")
    frame4 = ttkb.LabelFrame(lvl2_4root, bootstyle=SUCCESS, text="NALOG U WORD ")
    frame4.grid(row=0, column=0,padx=10, pady=10, sticky="news")
    frame4_1 = ttkb.LabelFrame(lvl2_4root, bootstyle=SUCCESS, text="NAPOMENA: ")
    frame4_1.grid(row=0,column=1, rowspan=2, padx=10, pady=10, sticky="news")
    label1 = ttkb.Label(frame4, text="BROJ NALOGA: \nILI OPSEG: ")
    label1.grid(row=0, column=0, padx=5, pady=10, sticky=E)
    label2 = ttkb.Label(frame4_1, text="U polje 'BROJ NALOGA' možete upisati broj \npojedinačnog naloga (e.g. 1982), \nili opseg koji uključuje \nobe granične vrednosti (e.g. 1982-1985), \nna ovaj način generisaće se word za \n1982, 1983, 1984 i 1985")
    label2.grid(row=0, column=2, rowspan=2)
    nalog_opseg = ttkb.Entry(frame4, width=15)
    nalog_opseg.grid(row=0, column=1, padx=5, pady=10)
    nalog_opseg.focus()

    submit_button = ttkb.Button(frame4, text="KREIRAJ NALOG", bootstyle=SUCCESS, width=15, command=on_submit)
    submit_button.grid(row=1, column=0, padx=5, pady=15)
    nalog_opseg.bind("<Return>", lambda event: on_submit())
    exit4_btn = ttkb.Button(frame4, text="ZATVORI", bootstyle=SUCCESS, width=15, command=close_toplevel_four)
    exit4_btn.grid(row=1, column=1, padx=5, pady=15)
    
# meni za sortiranje aktivnosti po lokacijama klijenta za izabrani period
def menu_3():
    global  lvl2_3root, start_date, end_date, sek_entry
    lvl2_3root = ttkb.Toplevel()
    lvl2_3root.title("IZVEŠTAJ ZA PERIOD")
    frame4 = ttkb.LabelFrame(lvl2_3root, bootstyle=SUCCESS)
    frame4.grid(row=0, column=0, padx=3, pady=3, sticky="news")  
    # etikete
    od_label = ttkb.Label(frame4, text="DATUM OD", bootstyle=SUCCESS, width=15)
    od_label.grid(row=1,column=0, sticky="w", pady=3, padx=3)
    do_label = ttkb.Label(frame4, text="DATUM DO", bootstyle=SUCCESS, width=15)
    do_label.grid(row=2, column=0, sticky="w", pady=3, padx=3)
    sek_label = ttkb.Label(frame4, text="SEKTOR:", bootstyle="success", width=15)
    sek_label.grid(row=0, column=0, sticky="w", pady=3, padx=3)
    # deo za unos
    start_date = ttkb.DateEntry(frame4, width=15, bootstyle="success", dateformat="%d.%m.%Y")
    start_date.grid(row=1, column=1, padx=8, pady=3, sticky="w")
    end_date = ttkb.DateEntry(frame4, width=15, bootstyle="success", dateformat="%d.%m.%Y")
    end_date.grid(row=2, column=1, padx=8, pady=3, sticky="w")
    sek_entry = ttkb.Combobox(frame4, width=15, bootstyle="success", values=["KT", "SPI", "SV"])
    sek_entry.grid(row=0, column=1, sticky="w", pady=3, padx=8)
    # dugmići
    kreiraj = ttkb.Button(frame4, text="KREIRAJ IZVEŠTAJ", width=15, bootstyle=SUCCESS, command=sta_je_radjeno)    
    kreiraj.grid(row=3, column=0, pady=15, padx=4)
    izađi = ttkb.Button(frame4, text="IZAĐI", width=15, bootstyle=SUCCESS, command=close_toplevel_tree)
    izađi.grid(row=3, column=1, pady=15, padx=4)

#  meni za mesečni izveštaj
def menu_2():
    global mesec_combo, lvl2root, search_activity_Entry
    lvl2root = ttkb.Toplevel()
    lvl2root.title("MESEČNI IZVEŠTAJ")
    frame2 = ttkb.LabelFrame(lvl2root, bootstyle=SUCCESS)
    frame2.grid(row=0, column=0, padx=10, pady=10, sticky="news")
    mesec_label = ttkb.Label(frame2, text="IZABERITE MESEC", bootstyle=SUCCESS)
    mesec_label.grid(row=0, column=0, padx=10, sticky="news")
    mesec_combo = ttkb.Combobox(frame2, values=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12], width=15, bootstyle=SUCCESS)
    mesec_combo.grid(row=1, column=0, padx=10, pady=10, sticky="news")
    search_activity_label = ttkb.Label(frame2, text="AKTIVNOST", bootstyle=SUCCESS)
    search_activity_label.grid(row=2, column=0, padx=10, sticky="news")
    search_activity_Entry = ttkb.Entry(frame2, bootstyle=SUCCESS)
    search_activity_Entry.grid(row=3, column=0, padx=10, sticky="news", columnspan=2)
    confirm_btn = ttkb.Button(frame2, text="POTVRDI", bootstyle=SUCCESS, width=15, command=on_button_click)
    confirm_btn.grid(row=4, column=0, padx=10, pady=10, sticky="news")
    exitlvl2_btn = ttkb.Button(frame2, text="IZAĐI", bootstyle="success", width=15, command=close_toplevel)
    exitlvl2_btn.grid(row=4, column=1, padx=10, pady=10, sticky="news")
# pravi izveštaj za period (poseban excel u kojem su sve adrese i šta je na tim adresama rađeno)
def sta_je_radjeno():
    # Connect to the SQLite3 database
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()

    # Get the dates from the DateEntry widgets
    od = start_date.entry.get()
    do = end_date.entry.get()
    sektor = sek_entry.get()
    # Parse the dates into a format SQLite can understand (YYYY-MM-DD)
    try:
        od = parser.parse(od, dayfirst=True).strftime("%Y-%m-%d")
        do = parser.parse(do, dayfirst=True).strftime("%Y-%m-%d")
    except ValueError:
        print("Invalid date format")
        return

    # Debug output to check date parsing
    print(f"Date Issued: {od}, Deadline: {do}")

    # Define the SQL query with reformatting of rok
    
    query = f"""
    SELECT klijent, mesto, adresa, aktivnost, rok
    FROM {table_name}
    WHERE 
        date(substr(rok, 7, 4) || '-' || substr(rok, 4, 2) || '-' || substr(rok, 1, 2)) >= ? 
        AND date(substr(rok, 7, 4) || '-' || substr(rok, 4, 2) || '-' || substr(rok, 1, 2)) <= ?
        AND sektor = ?
    """
    
    # Execute the query with the date parameters
    data = pd.read_sql_query(query, conn, params=[od, do, sektor])

    # Check if data is returned
    if data.empty:
        print("No data found for the given date range.")
    else:
        print(f"Data found: {len(data)} rows")

    # Ensure the date format is consistent in the DataFrame
    data['rok'] = pd.to_datetime(data['rok'], format="%d.%m.%Y", dayfirst=True)
    data['rok'] = data['rok'].dt.strftime("%d.%m.%Y")

    data['aktivnost'] = data['aktivnost'].apply(shorten_activity)

    # Create a pivot table to group activities by 'klijent', 'mesto', and 'Adresa'
    data['ActivityNum'] = data.groupby(['rok', 'klijent', 'mesto', 'adresa']).cumcount() + 1
    pivot_data = data.pivot(index=['rok', 'klijent', 'mesto', 'adresa'], columns='ActivityNum', values='aktivnost')

    # Rename the columns to reflect Aktivnost1, Aktivnost2, ...
    pivot_data.columns = [f"aktivnost{i}" for i in pivot_data.columns]

    # Reset the index to get 'Klijent', 'Mesto', and 'Adresa' as columns
    final_data = pivot_data.reset_index()

    # Save the data to an Excel file
    folder_path1 = r"PERIOD"
    report_filename = f'{folder_path1}\\izveštaj za period od {parser.parse(od).strftime("%d.%m.%Y")} do {parser.parse(do).strftime("%d.%m.%Y")}_{sektor}.xlsx'
    final_data.to_excel(report_filename, index=False)
    messagebox.showinfo(title="OBAVEŠTENJE", message=f"izveštaj za period od {parser.parse(od).strftime("%d.%m.%Y")} do {parser.parse(do).strftime("%d.%m.%Y")}_{sektor}.xlsx" + " je uspešno formiran")
    conn.close()

def fetch_imena():
    global imena
    veza = sq.connect("OSOBLJE.db")
    k = veza.cursor()
    k.execute("SELECT ime_i_prezime FROM kontrolori ORDER BY ime_i_prezime")
    imena = k.fetchall()
    values = [row[0] for row in imena]
    veza.close()

    return values

def fetch_firma():
    global firma
    veza = sq.connect("Evidencija.db")
    k = veza.cursor()
    k.execute("SELECT NAZIV FROM FIRME")
    firma = k.fetchall()
    values = [row[0] for row in firma]
    veza.close()

    return values

def populate_combobox(combobox, values):
    combobox['values'] = values

def unesi():
    global z, maxid
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()

    headers = f'''CREATE TABLE IF NOT EXISTS {table_name} 
                  (br_naloga INT, sektor TEXT, nalog_izdao TEXT, vođa_tima TEXT, 
                   član_1 TEXT, član_2 TEXT, član_3 TEXT, član_4 TEXT, 
                   klijent TEXT, mesto TEXT, adresa TEXT, aktivnost TEXT, 
                   datum_izdavanja DATE, rok DATE, kontakt_osoba TEXT, 
                   broj_dokumenta TEXT, invoice TEXT, napomena TEXT, rowid INT)'''
    c.execute(headers)

    query = f'''INSERT INTO {table_name} (br_naloga, sektor, nalog_izdao, vođa_tima, 
                  član_1, član_2, član_3, član_4, klijent, mesto, adresa, aktivnost, 
                  datum_izdavanja, rok, kontakt_osoba, broj_dokumenta, invoice, napomena, rowid) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)'''
    
    data = (num_entry.get(), combo.get(), issuer_entry.get(), reciver_entry.get(), 
            team_memeber_entry1.get(), team_memeber_entry2.get(), team_memeber_entry3.get(), team_memeber_entry4.get(), 
            client_entry.get(), city_entry.get(), adress_entry.get(), task_entry.get(), 
            date_issued_entry.entry.get(), deadline_entry.entry.get(), con_person_entry.get(), 
            task_No_entry.get(), invoice_no_entry.get(), note_entry.get(), rowid_entry.get())

    c.execute(query, data)
    conn.commit()
    

    messagebox.showinfo("OBAVEŠTENJE", f"Nalog broj {task_No_entry.get()} je dodat u {table_name}")

        
    values = (num_entry.get(), combo.get(), issuer_entry.get(), reciver_entry.get(), 
              team_memeber_entry1.get(), team_memeber_entry2.get(), team_memeber_entry3.get(), team_memeber_entry4.get(), 
              client_entry.get(), city_entry.get(), adress_entry.get(), task_entry.get(), date_issued_entry.entry.get(), 
              deadline_entry.entry.get(), con_person_entry.get(), task_No_entry.get(), invoice_no_entry.get(), note_entry.get(), rowid_entry.get())
        
    tree.insert("", END, values=values)

    task_No_entry.delete(0, "end")
    task_No_entry.insert(0, num_entry.get() + "." + combo.get() + "." + str(z+1))
    z += 1

    c.execute(f"SELECT MAX(rowid) FROM {table_name}")
    maxid = c.fetchone()[0]
    rowid_entry.delete(0, "end")
    rowid_entry.insert(0, maxid+1)
    conn.close()

# popunjava izdavaoce naloga u odnosu na izabrani sektor
def sektor_changed(event):
   
    AB = "Anđelko Baskić"
    ML = "Milomir Lukić"
    MT = "Milan Trišić"
    MB = "Miodrag Brklje"
    MV = "Milica Vučićević"
    SVI = "Nemanja Mitrović, Lazar Lešnjak, Dragana Mudavdžić"
    task_No_entry.delete(0, "end")
    task_No_entry.insert(0, num_entry.get() + "." + combo.get() + "." + str(z))

    if combo.get() == "KT":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, AB + " - " + ML)
        
    elif combo.get() == "SPI":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, MB + ", " + MT)
    elif combo.get() == "LB":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, MB + ", " + MT)
    elif combo.get() == "CE":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, MV)
    elif combo.get() == "SSP":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, SVI)       
    elif combo.get() == "SV":
        issuer_entry.delete(0, "end")
        issuer_entry.insert(0, MT)

def fetch_data(br_naloga):
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    c.execute(f"SELECT * FROM {table_name} WHERE br_naloga = ?", (br_naloga,))
    rows = c.fetchall()
    conn.close()
    return rows
# Function to set the font for each run in the document using XML
def set_font(run, font_name="Aptos Narrow", font_size=12, color=None, bold=False, underline=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.underline = underline

def set_font1(run, font_name="Aptos Narrow", font_size=11):
    
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

def set_repeat_table_header(row):
            """ Set repeat table row on every page via Word's internal XML structure """
            tr = row._element  # Get the row's XML element
            tbl_header = OxmlElement('w:tblHeader')  # Create the repeat header XML element
            tr.get_or_add_trPr().append(tbl_header)

## Function to create a Word document with the fetched data
def create_word_nalog(data, headers, br_naloga):
    global t1  # 't1' is the 'sektor' column value
    for item in data:
        t1 = item[1]  # sektor
        t2 = item[2]  # nalog_izdao
        t3 = item[12]  # datum_izdavanja
        t4 = item[13]  # rok
        vođa_tima = item[3]  # vođa_tima
        članovi_tima = [item[4], item[5], item[6], item[7]]  # član_1 to član_4

    # Filter out empty team members
    članovi_tima_str = ", ".join([član for član in članovi_tima if član])

    # Construct the filename
    doc_filename = f"Nalog_{br_naloga}.{t1}, {vođa_tima}"
    if članovi_tima_str:
        doc_filename += f", {članovi_tima_str}"

    doc_filename = doc_filename.strip()
    if not doc_filename.endswith(".docx"):
        doc_filename += ".docx"


    # Create the document
    doc = Document()

    # Set document to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.7)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)

    # Add header and logo
    doc.add_picture(r"IMAGES\LOGO_TP.png", width=Cm(1.5), height=Cm(1.5))
    heading = doc.add_heading('NALOG ZA RAD', 0)

    table_info = doc.add_table(rows=1, cols=2)
    table_info.style = None
    cell_left = table_info.rows[0].cells[0]
    cell_right = table_info.rows[0].cells[1]
    
    # Left column: Formatted general info
    info_paragraphs = [
        ("BROJ NALOGA: ", str(br_naloga)),
        ("SEKTOR: ", t1),
        ("NALOG IZDALI: ", t2 + "\n(K - 1; K - 2; K - 4, K - 6) - (K - 3; K - 5; K - 6; K - 7)" if t1 == "KT" else t2),
        ("DATUM IZDAVANJA / ROK: ", f"{t3} - {t4}")
    ]


    for label, value in info_paragraphs:
        p = cell_left.add_paragraph()
        run_label = p.add_run(label)
        set_font(run_label, "Aptos Narrow", 12, bold=True)  
        p.paragraph_format.space_after = Pt(0)
        run_value = p.add_run(value)
        set_font(run_value, "Aptos Narrow", 12, color=boja, underline=True)  

    # Right column: Activities List
    activities = [
        "K - 1: Kontrolisanje instalacija hidrantske mreže za gašenje požara",
        "K - 2: Kontrolisanje mobilnih uređaja za gašenje požara",
        "K - 3: Kontrolisanje instalacija i uređaja za automatsko otkrivanje i dojavu požara",
        "K - 4: Kontrolisanje instalacija i uređaja za gašenje požara",
        "K - 5: Kontrolisanje instalacija i uređaja za detekciju eksplozivnih i zapaljivih gasova",
        "K - 6: Kontrolisanje instalacija za odvođenje dima i toplote",
        "K - 7: Kontrolisanje instalacija i uređaja u zonama opasnosti od eksplozije"
    ]
    for activity in activities:
        p = cell_right.add_paragraph(activity)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        set_font(p.runs[0], "Aptos Narrow", 11)  

        # Add spacing before main data table
    dr = doc.add_paragraph("\n")
    dr.paragraph_format.space_after = Pt(0)

    included_headers = ['KLIJENT', 'MESTO', 'ADRESA', 'AKTIVNOST', 'KONTAKT', 'VOĐA \nTIMA', 'ČLANOVI \nTIMA', 'BR. DOK.', 'PONUDA \nUGOVOR', 'NAPOMENA']

    table = doc.add_table(rows=1, cols=len(included_headers), style="Light Grid Accent 5") 
    table.autofitContent = True
    table.allow_autofitContent = True

    hdr_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(included_headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            set_font(run, "Aptos Narrow", bold=True,  font_size=12)

# Loop over each row of data and populate the table
    for row in data:  # Ensure this loop iterates over every row
        row_cells = table.add_row().cells

# Adjusted row data, combining član_tima_1 to član_tima_4 while skipping empty values
        adjusted_row = [

            row[8],  # klijent
            row[9],  # mesto
            row[10], # adresa
            row[11], # aktivnost
            row[14], # kontakt_osoba
            row[3],  # vođa_tima

# Combine član_tima_1 to član_tima_4 into one cell, skipping empty values
            ', '.join([team_member for team_member in [row[4], row[5], row[6], row[7]] if team_member]),            
            row[15], # broj_dokumenta
            row[16], # invoice
            row[17]  # napomena
        ]
# Populate the table cells with the adjusted row data
        for i, item in enumerate(adjusted_row):
            row_cells[i].text = str(item) if item else ''
            for run in row_cells[i].paragraphs[0].runs:
                set_font(run, font_name="Aptos Narrow", font_size=11)  # Applying the font settings to each cell

# Save the document with a cleaner file name, skipping empty values
        team_members_str = ', '.join([team_member for team_member in [row[4], row[5], row[6], row[7]] if team_member])

        doc_filename = f"Nalog_{br_naloga}.{t1}, {row[3]}"

# If there are any team members, add them to the file name
        if team_members_str:
            doc_filename += f", {team_members_str}"

# Final file name with .docx extension
# Define the file path
    doc_filepath = r"Nalozi"

# Combine path and filename
    full_path = f"{doc_filepath}\\{doc_filename}"

# Ensure the file has the correct .docx extension
    if not full_path.endswith(".docx"):
        full_path += ".docx"  # Add .docx if missing

# Debug: Print the final path for confirmation
#    print(f"Final save path with extension: {full_path}")

    doc.save(full_path)
    messagebox.showinfo("ČUVANJE", f"Dokument '{doc_filename}' je uspešno kreiran.")

# handles creation of Nalog
def on_submit():
    input_value = nalog_opseg.get().strip()  # Get the input value and strip spaces
    if "-" in input_value:  # Check if input specifies a range
        try:
            start, end = map(int, input_value.split("-"))
            if start > end:
                raise ValueError("Invalid range: start must be less than or equal to end.")
        except ValueError as e:
            messagebox.showerror("Error", f"Invalid range: {e}")
            return

        for br_naloga in range(start, end + 1):  # Iterate over the range
            data = fetch_data(br_naloga)
            if data:
                headers = ['br_naloga', 'sektor', 'nalog_izdao', 'vođa_tima', 'član_1', 'član_2', 'član_3', 'član_4', 'klijent', 'mesto', 'adresa', 'aktivnost', 'datum_izdavanja', 'kontakt_osoba', 'broj_dokumenta', 'invoice', 'napomena', 'rowid']
                create_word_nalog(data, headers, br_naloga)
            else:
                print(f"No data found for 'br_naloga': {br_naloga}")
    else:  # Single number input
        try:
            br_naloga = int(input_value)
            data = fetch_data(br_naloga)
            if data:
                headers = ['br_naloga', 'sektor', 'nalog_izdao', 'vođa_tima', 'član_1', 'član_2', 'član_3', 'član_4', 'klijent', 'mesto', 'adresa', 'aktivnost', 'datum_izdavanja', 'kontakt_osoba', 'broj_dokumenta', 'invoice', 'napomena', 'rowid']
                create_word_nalog(data, headers, br_naloga)
            else:
                messagebox.showerror("Error", f"No data found for 'br_naloga': {br_naloga}")
        except ValueError:
            messagebox.showerror("Error", "Invalid input. Please enter a number or a range (e.g., 1792-1797).")

def close_app():
    if messagebox.askyesno("ZATVARANJE", "Da li želite da zatvorite aplikaciju?"):
        root.destroy()

def close_toplevel():
    if messagebox.askyesno("ZATVORI", "Da li želite da zatvorite prozor?"):
        lvl2root.destroy()

def close_toplevel_four():
    if messagebox.askyesno("ZATVORI", "Da li želite da zatvorite prozor?"):
        lvl2_4root.destroy()

def close_toplevel_tree():
    if messagebox.askyesno("ZATVORI", "Da li želite da zatvorite prozor?"):
        lvl2_3root.destroy()

def close_toplevel_two():
    if messagebox.askyesno("ZATVORI", "Da li želite da zatvorite prozor?"):
        lvl2_2root.destroy()

def filter_names(event, combo):
    # Get the typed text from the combobox
    typed_text = combo.get()

    # Filter the names list based on the typed text
    filtered_names = [name for name in imena if typed_text.lower() in name.lower()]

    # Update the combobox dropdown list with the filtered names
    combo['values'] = filtered_names

def reset():
    global z
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    c.execute(f"SELECT MAX(br_naloga) FROM {table_name}")
    maxnalog = c.fetchone()[0]
    today = datetime.date.today()
    formated_today = today.strftime("%d.%m.%Y")
       
    num_entry.delete(0, "end")
    if maxnalog == 0 or maxnalog == None:
        num_entry.insert(0, 1)
    else:
        num_entry.insert(0, int(maxnalog) + 1)
    issuer_entry.delete(0, "end")
    reciver_entry.delete(0, "end")
    team_memeber_entry1.delete(0, "end")
    team_memeber_entry2.delete(0, "end")
    team_memeber_entry3.delete(0, "end")
    team_memeber_entry4.delete(0, "end")
    con_person_entry.delete(0, "end")
    task_entry.delete(0, "end")
    client_entry.delete(0, "end")
    city_entry.delete(0, "end")
    adress_entry.delete(0, "end")
    note_entry.delete(0, "end")
    invoice_no_entry.delete(0, "end")
    combo.delete(0, "end")
    task_No_entry.delete(0, "end")
    date_issued_entry.entry.delete(0, "end")
    date_issued_entry.entry.insert(0, formated_today)
    deadline_entry.entry.delete(0, "end")
    deadline_entry.entry.insert(0, formated_today)
    c.execute(f"SELECT MAX(rowid) FROM {table_name}")
    maxid = c.fetchone()[0]
    rowid_entry.delete(0, "end")
    if maxid == 0 or maxid == None:
        rowid_entry.insert(0, 1)
    else:
        rowid_entry.insert(0, int(maxid) + 1)    
    z = 1

def edit():
    selected = tree.focus()
    values = (num_entry.get(), combo.get(), issuer_entry.get(), reciver_entry.get(), 
              team_memeber_entry1.get(), team_memeber_entry2.get(), team_memeber_entry3.get(), team_memeber_entry4.get(), 
              client_entry.get(), city_entry.get(), adress_entry.get(), task_entry.get(), date_issued_entry.entry.get(), 
              deadline_entry.entry.get(), con_person_entry.get(), task_No_entry.get(), invoice_no_entry.get(), note_entry.get(), rowid_entry.get())

    tree.item(selected, values=values)

    # Update the database
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    c.execute(f"""UPDATE {table_name} SET
        br_naloga = :nalog,
        sektor = :sektor,
        nalog_izdao = :izdavalac,
        vođa_tima = :vođa,
        član_1 = :član_1,
        član_2 = :član_2,
        član_3 = :član_3,
        član_4 = :član_4,
        klijent = :klijent,
        mesto = :mesto,
        adresa = :adresa,
        aktivnost = :aktivnost,
        datum_izdavanja = :dizd,
        rok = :rok,
        kontakt_osoba = :kontakt,
        broj_dokumenta = :isprava,
        invoice = :ponuda,
        napomena = :note
        
        WHERE rowid = :rowid""",
        {
            "nalog" : num_entry.get(),
            "sektor" : combo.get(),
            "izdavalac" : issuer_entry.get(), 
            "vođa" : reciver_entry.get(), 
            "član_1" : team_memeber_entry1.get(), 
            "član_2" : team_memeber_entry2.get(), 
            "član_3" : team_memeber_entry3.get(), 
            "član_4" : team_memeber_entry4.get(), 
            "klijent" : client_entry.get(), 
            "mesto" : city_entry.get(), 
            "adresa" : adress_entry.get(), 
            "aktivnost" : task_entry.get(), 
            "dizd" : date_issued_entry.entry.get(), 
            "rok" : deadline_entry.entry.get(), 
            "kontakt" : con_person_entry.get(), 
            "isprava" : task_No_entry.get(), 
            "ponuda" : invoice_no_entry.get(), 
            "note" : note_entry.get(),
            "rowid" : rowid_entry.get()

        })
    
    conn.commit()
    conn.close()

def edit2():
    selected1 = drvo.focus()
    values1 = (entry_klijent.get(), entry_ponuda.get(),  entry_kontakt.get(), entry_mesto.get(), entry_adresa.get(), entry_pun_naziv.get())
    drvo.item(selected1, values=values1)
    
    # Update the database
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    c.execute("""UPDATE klijent SET
        firma = :firma,
        mesto = :mesto,
        adresa = :adresa,
        "BROJ PONUDE / UGOVORA" = :ponuda,
        kontakt = :kontakt,
        "PUN NAZIV" = :punnaziv
        
        WHERE firma = :firma""",
            {
            "firma" : entry_klijent.get(),
            "mesto" : entry_mesto.get(),
	        "adresa" : entry_adresa.get(),
	        "ponuda" : entry_ponuda.get(),
	        "kontakt" : entry_kontakt.get(),
	        "punnaziv" : entry_pun_naziv.get()

            })
    
    conn.commit()
    conn.close()

def select_record(event):
    
    num_entry.delete(0, END)
    combo.delete(0, END)
    issuer_entry.delete(0, END)
    reciver_entry.delete(0, END)
    team_memeber_entry1.delete(0, END)
    team_memeber_entry2.delete(0, END)
    team_memeber_entry3.delete(0, END)
    team_memeber_entry4.delete(0, END)
    client_entry.delete(0, END)
    city_entry.delete(0, END)
    adress_entry.delete(0, END)
    task_entry.delete(0, END)
    date_issued_entry.entry.delete(0, END)
    deadline_entry.entry.delete(0, END)
    con_person_entry.delete(0, END)
    task_No_entry.delete(0, END)
    invoice_no_entry.delete(0, END)
    note_entry.delete(0, END)
    rowid_entry.delete(0, END)

    selected = tree.focus()
	
    values = tree.item(selected, 'values')

    num_entry.insert(0, values[0])
    combo.insert(0, values[1])
    issuer_entry.insert(0, values[2])
    reciver_entry.insert(0, values[3])
    team_memeber_entry1.insert(0, values[4])
    team_memeber_entry2.insert(0, values[5])
    team_memeber_entry3.insert(0, values[6])
    team_memeber_entry4.insert(0, values[7])
    client_entry.insert(0, values[8])
    city_entry.insert(0, values[9])
    adress_entry.insert(0, values[10])
    task_entry.insert(0, values[11])
    date_issued_entry.entry.insert(0, values[12])
    deadline_entry.entry.insert(0, values[13])
    con_person_entry.insert(0, values[14])
    task_No_entry.insert(0, values[15])
    invoice_no_entry.insert(0, values[16])
    note_entry.insert(0, values[17])
    rowid_entry.insert(0, values[18])

def select_record_in_klijent(event):
    
    entry_klijent.delete(0, END)
    entry_mesto.delete(0, END)
    entry_adresa.delete(0, END)
    entry_ponuda.delete(0, END)
    entry_kontakt.delete(0, END)
    entry_pun_naziv.delete(0, END)
    identry.delete(0, END)

    selected1 = drvo.focus()
    values1 = drvo.item(selected1, 'values')

    entry_klijent.insert(0, values1[0])
    entry_mesto.insert(0, values1[3])
    entry_adresa.insert (0, values1[4])
    entry_ponuda.insert (0, values1[1])
    entry_kontakt.insert (0, values1[2])
    entry_pun_naziv.insert (0, values1[5])
    identry.insert(0, values1[6])

def removeOneKlijent():
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    idrow = identry.get()
    if messagebox.askyesno("UPOZORENJE", "Da li zaista želite da izbrišete selektovani unos?"):
        c.execute("DELETE FROM klijent WHERE IDrow = ?", (idrow,))
    # Remove the selected item from the treeview
        x = drvo.selection()[0]
        drvo.delete(x)

        conn.commit()
        conn.close()

def removeOne():
    table_name = get_table_name()
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()

    rowid_value = rowid_entry.get()  
    if messagebox.askyesno("UPOZORENJE", "Da li zaista želite da izbrišete selektovani unos?"):
        c.execute(f"DELETE FROM {table_name} WHERE rowid = ?", (rowid_value,))
    # Remove the selected item from the treeview
        x = tree.selection()[0]
        tree.delete(x)

        conn.commit()
        conn.close()

def newid():
    table_name = get_table_name()
    c.execute(f"SELECT MAX(rowid) FROM {table_name}")
    maxid = c.fetchone()[0]
    rowid_entry.delete(0, "end")
    if maxid == 0 or maxid == None:
        rowid_entry.insert(0, 1)
    else:
        rowid_entry.insert(0, int(maxid) + 1)

def k1():
    k_1 = "Periodično kontrolisanje instalacija hidrantske mreže za gašenje požara"
    task_entry.delete(0, END)
    task_entry.insert(0, k_1)

def k2():
    k_2 = "Periodično kontrolisanje mobilnih uređaja za gašenje požara"
    task_entry.delete(0, END)
    task_entry.insert(0, k_2)

def k3():
    k_3 = "Periodično kontrolisanje instalacija i uređaja za automatsko otkrivanje i dojavu požara"
    task_entry.delete(0, END)
    task_entry.insert(0, k_3)

def k4():
    k_4 = "Periodično kontrolisanje instalacija i uređaja za gašenje požara"
    task_entry.delete(0, END)
    task_entry.insert(0, k_4)

def k5():
    k_5 = "Periodično kontrolisanje instalacija i uređaja za detekciju eksplozivnih i zapaljivih gasova"
    task_entry.delete(0, END)
    task_entry.insert(0, k_5)

def k6():
    k_6 = "Periodično kontrolisanje instalacija za odvođenje dima i toplote"
    task_entry.delete(0, END)
    task_entry.insert(0, k_6)

def k7():
    k_7 = "Periodično kontrolisanje instalacija i uređaja u zonama opasnosti od eksplozije"
    task_entry.delete(0, END)
    task_entry.insert(0, k_7)

def im1():
    im_1 = "Ispitivanje mikroklime u zimskom periodu i osvetljenosti"
    task_entry.delete(0, END)
    task_entry.insert(0, im_1)

def im2():
    im_2 = "Ispitivanje mikroklime u letnjem periodu i osvetljenosti"
    task_entry.delete(0, END)
    task_entry.insert(0, im_2)

def im3():
    im_3 = "Pregled i ispitivanje gromobranskih instalacija"
    task_entry.delete(0, END)
    task_entry.insert(0, im_3)

def im4():
    im_4 = "Pregled i ispitivanje električnih instalacija"
    task_entry.delete(0, END)
    task_entry.insert(0, im_4)

def im5():
    im_5 = "Kontrolisanje PP rasvete"
    task_entry.delete(0, END)
    task_entry.insert(0, im_5)

def im6():
    im_6 = "Pregled opreme"
    task_entry.delete(0, END)
    task_entry.insert(0, im_6)

def im7():
    im_7 = "Merenje i ispitivanje hemijskih štetnosti"
    task_entry.delete(0, END)
    task_entry.insert(0, im_7)

def lb1():
    lb1 = "Pregled i ispitivanje izolacionih aparata"
    task_entry.delete(0, END)
    task_entry.insert(0, lb1)

def lb2():
    lb2 = "Pregled i ispitivanje creva"
    task_entry.delete(0, END)
    task_entry.insert(0, lb2)

def lb3():
    lb3 = "Pregled i ispitivanje statičkog elektriciteta"
    task_entry.delete(0, END)
    task_entry.insert(0, lb3)

def new_client():
    global entry_klijent, lvl2_2root, entry_adresa, entry_kontakt, entry_mesto, entry_ponuda, entry_pun_naziv, drvo, identry
    lvl2_2root = ttkb.Toplevel()
    lvl2_2root.title("BAZA KLIJENATA")
    frame3 = ttkb.LabelFrame(lvl2_2root)
    frame3.grid(row=0, column=1)
    frame2 = ttkb.LabelFrame(lvl2_2root, bootstyle=SUCCESS)
    frame2.grid(row=0, column=0, padx=10, pady=10, sticky="news")
    label_klijent = ttkb.Label(frame2, text="NAZIV KLIJENTA", bootstyle=SUCCESS)
    label_klijent.grid(row=0, column=0, padx=10, sticky="news", pady=10)
    label_mesto = ttkb.Label(frame2, text="MESTO", bootstyle=SUCCESS)
    label_mesto.grid(row=1, column=0, padx=10, pady=10, sticky="news")
    label_adresa = ttkb.Label(frame2, text="ADRESA", bootstyle=SUCCESS)
    label_adresa.grid(row=2, column=0, padx=10, sticky="news", pady=10)
    label_ponuda = ttkb.Label(frame2, text="PONUDA / UGOVOR", bootstyle=SUCCESS)
    label_ponuda.grid(row=3, column=0, padx=10, sticky="news", pady=10)
    label_kontakt = ttkb.Label(frame2, text="KONTAKT OSOBA", bootstyle=SUCCESS)
    label_kontakt.grid(row=4, column=0, padx=10, sticky="news", pady=10)
    label_pun_naziv = ttkb.Label(frame2, text="NAZIV, MESTO, ADRESA", bootstyle=SUCCESS)
    label_pun_naziv.grid(row=5, column=0, padx=10, sticky="news", pady=10)
    idrow_label = ttkb.Label(frame2, text="ID KLIJENTA", bootstyle=SUCCESS)
    idrow_label.grid(row=6, column=0, padx=10, sticky="news", pady=10)

    entry_klijent = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_klijent.grid(row=0, column=1, padx=10, sticky="news", pady=10, columnspan=2)
    entry_mesto = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_mesto.grid(row=1, column=1, padx=10, pady=10, sticky="news", columnspan=2)
    entry_adresa = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_adresa.grid(row=2, column=1, padx=10, sticky="news", pady=10, columnspan=2)
    entry_ponuda = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_ponuda.grid(row=3, column=1, padx=10, sticky="news", pady=10, columnspan=2)
    entry_kontakt = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_kontakt.grid(row=4, column=1, padx=10, sticky="news", pady=10, columnspan=2)
    entry_pun_naziv = ttkb.Entry(frame2, bootstyle=SUCCESS, width=50)
    entry_pun_naziv.grid(row=5, column=1, padx=10, sticky="news", pady=10, columnspan=2)
    identry = ttkb.Entry(frame2, bootstyle=SUCCESS)
    identry.grid(row=6, column=1, padx=10, sticky="news", pady=10, columnspan=2)

    confirm_btn = ttkb.Button(frame2, text="POTVRDI", bootstyle=SUCCESS, width=20, command=zapamti_klijenta)
    confirm_btn.grid(row=7, column=0, padx=10, pady=15)
    exitlvl2_btn = ttkb.Button(frame2, text="IZAĐI", bootstyle="success", width=20, command=close_toplevel_two)
    exitlvl2_btn.grid(row=7, column=2, padx=10, pady=15)
    editlvl2_btn = ttkb.Button(frame2, text="IZMENI", bootstyle="success", width=20, command=edit2) 
    editlvl2_btn.grid(row=7, column=1, padx=10, pady=15)
    brišiunos_btn = ttkb.Button(frame2, text="IZBRIŠI UNOS", bootstyle=SUCCESS, width=20, command=removeOneKlijent)
    brišiunos_btn.grid(row=8, column=0, padx=10, pady=15)
    # Add Drvoview widget
    drvo = ttkb.Treeview(frame3, columns=("firma", "ponuda", "kontakt", "mesto", "adresa", "pun_naziv"), show="headings", height=23)
    drvo.grid(row=0, column=0, sticky="news")

    # Define headings
    drvo.heading("firma", text="Firma")
    drvo.heading("ponuda", text="Ponuda / Ugovor")
    drvo.heading("kontakt", text="Kontakt")
    drvo.heading("mesto", text="Mesto")
    drvo.heading("adresa", text="Adresa")
    drvo.heading("pun_naziv", text="Pun Naziv")

    # Populate Drvoview with data
    populate_drvoview()
    drvo.bind("<ButtonRelease-1>", select_record_in_klijent)
    # Scrollbar for Drvoview
    drvo_scroll = ttkb.Scrollbar(frame3, orient="vertical", command=drvo.yview)
    drvo_scroll.grid(row=0, column=1, sticky="ns")
    drvo.configure(yscrollcommand=drvo_scroll.set)

def populate_drvoview():
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()

    # Fetch data from the 'klijent' table
    c.execute("SELECT * FROM klijent ORDER BY FIRMA ASC")
    rows = c.fetchall()

    # Insert each row into the drvoview
    for row in rows:
        drvo.insert("", "end", values=row)

    conn.close()
   
def zapamti_klijenta():
    naziv = entry_klijent.get()
    ponuda = entry_ponuda.get()
    kontakt = entry_kontakt.get()
    mesto = entry_mesto.get()
    adresa = entry_adresa.get()
    pun_naziv = entry_klijent.get() + ", " + entry_mesto.get() + ", " + entry_adresa.get()

    conn = sq.connect("Evidencija.db")
    c = conn.cursor()

    data_insert_query = '''INSERT INTO klijent ("FIRMA", "BROJ PONUDE / UGOVORA", "KONTAKT", "MESTO", "ADRESA", "PUN NAZIV" 
                    ) VALUES 
                (?, ?, ?, ?, ?, ?)'''

    data_insert_tuple = (naziv, ponuda, kontakt,mesto, adresa, pun_naziv)
        
    
    c.execute(data_insert_query, data_insert_tuple)
    conn.commit()
    messagebox.showinfo(title="OBAVEŠTENJE", message=r"Klijent: " + entry_klijent.get() + " je unet u bazu klijenata!!!")    
    entry_klijent.delete(0, "end")
    entry_ponuda.delete(0, "end")
    entry_kontakt.delete(0, "end")
    entry_mesto.delete(0, "end")
    entry_adresa.delete(0, "end")

def update_widgets(event):
# updates data on client and populates appropriate widgets    
    # Get the partial text entered in the "FIRMA" entry widget
    partial_firma = firma_var.get()
    
    # Query the database for matching records
    conn = sq.connect('Evidencija.db')
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT * FROM klijent WHERE FIRMA LIKE ?
    """, (partial_firma + '%',))
    
    result = cursor.fetchone()
    conn.close()
    
    if result:
        # Fill the other entry widgets with data from the selected row
        ponuda_var.set(result[1])
        kontakt_var.set(result[2])
        mesto_var.set(result[3])
        adresa_var.set(result[4])
        pun_naziv_var.set(result[5])

def fill_firma_on_focus_out(event):
    if pun_naziv_var.get():
        firma_var.set(pun_naziv_var.get())

#all frames
basic_frame =ttkb.LabelFrame(root, text="OSNOVNI PODACI", bootstyle="success.TLabelframe")
basic_frame.grid(row=0, column=0, padx=5, pady=5, sticky="news")
client_frame =ttkb.LabelFrame(root, text="PODACI O KLIJENTU", bootstyle="success")
client_frame.grid(row=1, column=0, padx=5, pady=5, sticky="news")
btns1_frame =ttkb.LabelFrame(root, text="NALOG", bootstyle="success")
btns1_frame.grid(row=3, column=0, padx=5, pady=5, sticky="news")
btns2_frame =ttkb.LabelFrame(root, text="IZVEŠTAVANJE", bootstyle="success")
btns2_frame.grid(row=3, column=1, columnspan = 3, padx=5, pady=5, sticky="news")
task_frame =ttkb.LabelFrame(root, text="PODACI O ZADATKU", bootstyle="success")
task_frame.grid(row=0, column=1, columnspan=3, padx=5, pady=5, sticky="news")
kt_tasks_frame =ttkb.LabelFrame(root, text="KONTROLNO TELO", bootstyle="success")
kt_tasks_frame.grid(row=1, column=1, padx=5, pady=5, sticky="news")
spi_tasks_frame =ttkb.LabelFrame(root, text="SEKTOR PREGLEDA I ISPITIVANJA", bootstyle="success")
spi_tasks_frame.grid(row=1, column=2, padx=5, pady=5, sticky="news")
lb_tasks_frame =ttkb.LabelFrame(root, text="LABORATORIJA", bootstyle="success")
lb_tasks_frame.grid(row=1, column=3, padx=5, pady=5, sticky="news")
tree_frame =ttkb.Frame(root)
tree_frame.grid(row=4, column=0, columnspan = 4, padx=5, pady=5, sticky="news")

#all labels widgets

    #labels for basic frame
sec_label =ttkb.Label(basic_frame, text="Sektor:", width=15, font=("Ariel", 10)).grid(row=1, column=0)
num_label =ttkb.Label(basic_frame, text="Br. Naloga:", width=15, font=("Ariel", 10)).grid(row=2, column=0)
issuer_label =ttkb.Label(basic_frame, text="Izdavalac naloga:", width=15, font=("Ariel", 10)).grid(row=3, column=0)
reciver_label =ttkb.Label(basic_frame, text="Vođa tima:", width=15, font=("Ariel", 10)).grid(row=4, column=0)
team_member_label =ttkb.Label(basic_frame, text="Članovi tima: ", width=15, font=("Ariel", 10)).grid(row=5, column=0)
date_issued_label =ttkb.Label(basic_frame, text="Datum izdavanja naloga:", width=25, font=("Ariel", 10)).grid(row=0, column=3)
deadline_label =ttkb.Label(basic_frame, text="Rok za izvršenje naloga:", width=25, font=("Ariel", 10)).grid(row=1, column=3)
con_person_label =ttkb.Label(basic_frame, text="Osoba za kontakt:", width=25, font=("Ariel", 10)).grid(row=2, column=3)

    #labels for client frame
client_label =ttkb.Label(client_frame, text="Klijent", width=15, font=("Ariel", 10)).grid(row=0, column=0)
city_label =ttkb.Label(client_frame, text="Mesto:", width=15, font=("Ariel", 10)).grid(row=1, column=0)
addres_label =ttkb.Label(client_frame, text="Adresa:", width=15, font=("Ariel", 10)).grid(row=0, column=2)
pun_naziv_label =ttkb.Label(client_frame, text="Naziv i adresa:", width=15, font=("Ariel", 10)).grid(row=2, column=0)
invoice_no_label =ttkb.Label(client_frame, text="Br. Ponude", width=15, font=("Ariel", 10)).grid(row=1, column=2)
    
    #labels for task frame
task_No_label =ttkb.Label(task_frame, text="Br. zadatka:", width=25, font=("Ariel", 10)).grid(row=0, column=0, pady=5, sticky="w")
unique_id_label =ttkb.Label(task_frame, text="ID", width=5, font=("Ariel", 10)).grid(row=0, column=1, sticky="e")
task_label =ttkb.Label(task_frame, text="Aktivnost / Zadatak:", width=40, font=("Ariel", 10)).grid(row=2, column=0, pady=5, sticky="w")
note_label =ttkb.Label(task_frame, text="Napomena:", width=40, font=("Ariel", 10)).grid(row=4, column=0, pady=5, sticky="w")

#all Entry widgets
    #in basic frame
combo =ttkb.Combobox(basic_frame, values=["KT", "LB", "SPI", "SSP", "CE", "SV"], width=15, style="light")
combo.grid(row=1, column=1, padx=10,pady=10, sticky="w")
num_entry =ttkb.Entry(basic_frame, width=20, style="light")
num_entry.grid(row=2, column=1, padx=10,pady=10)

if maxnalog == 0 or maxnalog is None:
    num_entry.insert(0, 1)
else:
    num_entry.insert(0, int(maxnalog) + 1)

issuer_entry =ttkb.Entry(basic_frame, width=20, style="light")
issuer_entry.grid(row=3, column=1, padx=10,pady=10)
reciver_entry =ttkb.Combobox(basic_frame, width=20, style="light")
reciver_entry.grid(row=4, column=1, padx=10,pady=10)
team_memeber_entry1 =ttkb.Combobox(basic_frame, width=20, style="light")
team_memeber_entry1.grid(row=5, column=1, padx=10,pady=10)
team_memeber_entry2 =ttkb.Combobox(basic_frame, width=20, style="light")
team_memeber_entry2.grid(row=5, column=2, padx=10,pady=10)
team_memeber_entry3 =ttkb.Combobox(basic_frame, width=20, style="light")
team_memeber_entry3.grid(row=5, column=3, padx=10,pady=10)
team_memeber_entry4 =ttkb.Combobox(basic_frame, width=20, style="light")
team_memeber_entry4.grid(row=5, column=4, padx=10,pady=10)
date_issued_entry =ttkb.DateEntry(basic_frame, width=15, bootstyle="success", dateformat="%d.%m.%Y")
date_issued_entry.grid(row=0, column=4,pady=10)
deadline_entry =ttkb.DateEntry(basic_frame, width=15, bootstyle="success", dateformat="%d.%m.%Y")
deadline_entry.grid(row=1, column=4,pady=10)
con_person_entry =ttkb.Entry(basic_frame, width=20, style="light", textvariable=kontakt_var)
con_person_entry.grid(row=2, column=4, padx=10,pady=10)

    #in client frame
client_entry =ttkb.Entry(client_frame, width=20, style="light", textvariable=firma_var)
client_entry.grid(row=0, column=1, padx=10,pady=10, sticky="w")
city_entry =ttkb.Entry(client_frame, width=20, style="light", textvariable=mesto_var)
city_entry.grid(row=1, column=1, padx=10,pady=10, sticky="w")
adress_entry =ttkb.Entry(client_frame, width=30, style="light", textvariable=adresa_var)
adress_entry.grid(row=0, column=3, padx=10,pady=10, sticky="w")
invoice_no_entry =ttkb.Entry(client_frame, width=20, style="light", textvariable=ponuda_var)
invoice_no_entry.grid(row=1, column=3, padx=10,pady=10, sticky="w")
pun_naziv_entry =ttkb.Entry(client_frame, width=50, style="light", textvariable=pun_naziv_var)
pun_naziv_entry.grid(row=2, column=1, padx=10,pady=10, sticky="w")

    #in task frame
task_No_entry =ttkb.Entry(task_frame, width=20, style="light")
task_No_entry.grid(row=1, column=0, padx=5,pady=5, columnspan=2, sticky="w")
rowid_entry =ttkb.Entry(task_frame, width=5, style="light")
rowid_entry.grid(row=1, column=1, sticky="e")
c.execute("SELECT MAX(rowid) FROM Evidencija2025")
maxid = c.fetchone()[0]

if maxid == 0 or maxnalog is None or maxid is None:
    rowid_entry.insert(0, 1)
else:
    rowid_entry.insert(0, int(maxid) + 1)

task_entry =ttkb.Entry(task_frame, width=85, style="light")
task_entry.grid(row=3, column=0, padx=5,pady=5, columnspan=2, sticky="w")
note_entry =ttkb.Entry(task_frame, width=85, style="light")
note_entry.grid(row=5,column=0, padx=5,pady=5, columnspan=2, sticky="w")
set_id_btn =ttkb.Button(task_frame, text="RESET", width=5, style="light", command=newid)
set_id_btn.grid(row=1, column=2)

#all btns in btn frame
fill_btn =ttkb.Button(btns1_frame, text="UNESI NALOG", width=20, style="solid toolbutton", command=unesi)
fill_btn.grid(row=0, column=0, padx=20, pady=5, sticky="w")
reset_btn =ttkb.Button(btns1_frame, text="NOV NALOG", width=20, style="solid toolbutton", command=reset)
reset_btn.grid(row=0, column=1, padx=20, pady=5, sticky="we")
exit_btn =ttkb.Button(btns1_frame, text="ZATVORI", width=20, style="solid toolbutton", command=close_app)
exit_btn.grid(row=0, column=4, padx=20, pady=5, sticky="e")
remember_cl_btn =ttkb.Button(btns2_frame, text="EVIDENCIJA KLIJENATA", width=20, style="solid toolbutton",command=new_client)
remember_cl_btn.grid(row=0, column=0, padx=20, pady=5, sticky="w")
print_to_pdf_btn =ttkb.Button(btns1_frame, text="IZBRIŠI UNOS", width=20, style="outline toolbutton", command=removeOne)
print_to_pdf_btn.grid(row=1, column=0, padx=20, pady=5, sticky="we")
print_report_btn =ttkb.Button(btns2_frame, text="MESEČNI IZVEŠTAJ", width=20, style="solid toolbutton", command=menu_2)
print_report_btn.grid(row=0, column=1, padx=20, pady=5, sticky="we")
EDIT_btn =ttkb.Button(btns1_frame, text="IZMENI UNOS", width=20, style="solid toolbutton", command=edit)
EDIT_btn.grid(row=0, column=2, padx=20, pady=5, sticky="we")
period_rep =ttkb.Button(btns2_frame, text="IZVEŠTAJ ZA PERIOD", width=20, style="solid toolbutton", command=menu_3)
period_rep.grid(row=0, column=2, padx=20, pady=5, sticky="we")
nalog_u_doc =ttkb.Button(btns1_frame, text="NALOG U WORD", width=20, style="solid toolbutton", command=menu_4) 
nalog_u_doc.grid(row=0, column=3, padx=20, pady=5, sticky="we")
seticuse2_btn =ttkb.Button(btns2_frame, text="SAČUVAJ LOKACIJU", width=20, style="solid toolbutton", state="disabled")
seticuse2_btn.grid(row=1, column=0, padx=20, pady=5, sticky="we")

    #in kt frame
KT1 =ttkb.Button(kt_tasks_frame, text="K - 1", width=5, style="solid toolbutton", command=k1).grid(row=0, column=0, padx=8, pady=8)
KT2 =ttkb.Button(kt_tasks_frame, text="K - 2", width=5, style="solid toolbutton", command=k2).grid(row=0, column=1, padx=8, pady=8)
KT3 =ttkb.Button(kt_tasks_frame, text="K - 3", width=5, style="solid toolbutton", command=k3).grid(row=0, column=2, padx=8, pady=8)
KT4 =ttkb.Button(kt_tasks_frame, text="K - 4", width=5, style="solid toolbutton", command=k4).grid(row=1, column=0, padx=8, pady=8)
KT5 =ttkb.Button(kt_tasks_frame, text="K - 5", width=5, style="solid toolbutton", command=k5).grid(row=1, column=1, padx=8, pady=8)
KT6 =ttkb.Button(kt_tasks_frame, text="K - 6", width=5, style="solid toolbutton", command=k6).grid(row=1, column=2, padx=8, pady=8)
KT7 =ttkb.Button(kt_tasks_frame, text="K - 7", width=5, style="solid toolbutton", command=k7).grid(row=2, column=0, padx=8, pady=8)
KT8 =ttkb.Button(kt_tasks_frame, text="K - 8", width=5, style="solid toolbutton", state="disabled").grid(row=2, column=1, padx=8, pady=8)
KT9 =ttkb.Button(kt_tasks_frame, text="K - 9", width=5, style="solid toolbutton", state="disabled").grid(row=2, column=2, padx=8, pady=8)

    #in SPI frame
SPI1 =ttkb.Button(spi_tasks_frame, text="SPI - 1", width=6, style="solid toolbutton", command=im1).grid(row=0, column=0, padx=8, pady=8)
SPI2 =ttkb.Button(spi_tasks_frame, text="SPI - 2", width=6, style="solid toolbutton", command=im2).grid(row=0, column=1, padx=8, pady=8)
SPI3 =ttkb.Button(spi_tasks_frame, text="SPI - 3", width=6, style="solid toolbutton", command=im3).grid(row=0, column=2, padx=8, pady=8)
SPI4 =ttkb.Button(spi_tasks_frame, text="SPI - 4", width=6, style="solid toolbutton", command=im4).grid(row=1, column=0, padx=8, pady=8)
SPI5 =ttkb.Button(spi_tasks_frame, text="SPI - 5", width=6, style="solid toolbutton", command=im5).grid(row=1, column=1, padx=8, pady=8)
SPI6 =ttkb.Button(spi_tasks_frame, text="SPI - 6", width=6, style="solid toolbutton", command=im6).grid(row=1, column=2, padx=8, pady=8)
SPI7 =ttkb.Button(spi_tasks_frame, text="SPI - 4", width=6, style="solid toolbutton", command=im7).grid(row=2, column=0, padx=8, pady=8)
SPI5 =ttkb.Button(spi_tasks_frame, text="SPI - 5", width=6, style="solid toolbutton", state="disabled").grid(row=2, column=1, padx=8, pady=8)
SPI6 =ttkb.Button(spi_tasks_frame, text="SPI - 6", width=6, style="solid toolbutton", state="disabled").grid(row=2, column=2, padx=8, pady=8) 
   
   # in LB frame
LB1 =ttkb.Button(lb_tasks_frame, text="LB - 1", width=5, style="solid toolbutton", command=lb1).grid(row=0, column=0, padx=8, pady=8)
LB2 =ttkb.Button(lb_tasks_frame, text="LB - 2", width=5, style="solid toolbutton", command=lb2).grid(row=0, column=1, padx=8, pady=8)
LB3 =ttkb.Button(lb_tasks_frame, text="LB - 3", width=5, style="solid toolbutton", command=lb3).grid(row=0, column=2, padx=8, pady=8)
LB4 =ttkb.Button(lb_tasks_frame, text="LB - 4", width=5, style="solid toolbutton", state="disabled").grid(row=1, column=0, padx=8, pady=8)
LB5 =ttkb.Button(lb_tasks_frame, text="LB - 5", width=5, style="solid toolbutton", state="disabled").grid(row=1, column=1, padx=8, pady=8)
LB6 =ttkb.Button(lb_tasks_frame, text="LB - 6", width=5, style="solid toolbutton", state="disabled").grid(row=1, column=2, padx=8, pady=8)
LB7 =ttkb.Button(lb_tasks_frame, text="LB - 7", width=5, style="solid toolbutton", state="disabled").grid(row=2, column=0, padx=8, pady=8)
LB8 =ttkb.Button(lb_tasks_frame, text="LB - 5", width=5, style="solid toolbutton", state="disabled").grid(row=2, column=1, padx=8, pady=8)
LB9 =ttkb.Button(lb_tasks_frame, text="LB - 6", width=5, style="solid toolbutton", state="disabled").grid(row=2, column=2, padx=8, pady=8)
    
    #in tree frame
scroll = ttkb.Scrollbar(tree_frame, orient="vertical", style="success")
scroll.pack(side=RIGHT, fill=Y)
scrollx = ttkb.Scrollbar(tree_frame, orient="horizontal", style="success")
scrollx.pack(side=BOTTOM, fill=X)

# Create a treeview
tree = ttk.Treeview(tree_frame, yscrollcommand=scroll.set, xscrollcommand=scrollx.set, selectmode="browse", show="tree headings", height=200)
tree.pack()

scroll.config(command=tree.yview)
scrollx.config(command=tree.xview)

tree.tag_configure('oddrow', background="#0D222F")
tree.tag_configure('evenrow', background="#16394F")
style.map('Treeview', background=[('selected', "#DAA520")])

tree_frame.configure(width=1000, height=300)
basic_frame.configure(width=600, height=300)
client_frame.configure(width=400, height=300)

tree["columns"] = ("br_naloga", "sektor", "izdavalac", "vođa_tima", 
                    "član_tima_1", "član_tima_2", "član_tima_3", "član_tima_4", 
                    "klijent", "mesto", "adresa", "aktivnost", "datum_izdavanja", "rok", "kontakt", "broj_dokumenta", "br_ponude", "Napomena", "rowid")

tree.column("#0", width=0, stretch=NO)
tree.column("br_naloga", anchor=CENTER, width=60, stretch=True)
tree.column("sektor", anchor=CENTER, width=80, stretch=True)
tree.column("izdavalac", anchor=W, width=150, stretch=True)
tree.column("vođa_tima", anchor=W, width=150, stretch=True)
tree.column("član_tima_1", anchor=W, width=150, stretch=True)
tree.column("član_tima_2", anchor=W, width=150, stretch=True)
tree.column("član_tima_3", anchor=W, width=150, stretch=True)
tree.column("član_tima_4", anchor=W, width=150, stretch=True)
tree.column("klijent", anchor=W, width=150, stretch=True)
tree.column("mesto", anchor=W, width=120, stretch=True)
tree.column("adresa", anchor=W, width=150, stretch=True)
tree.column("aktivnost", anchor=W, width=400, stretch=True)
tree.column("datum_izdavanja", anchor=CENTER, width=110, stretch=True)
tree.column("rok", anchor=CENTER, width=110, stretch=True)
tree.column("kontakt", anchor=W, width=200, stretch=True)
tree.column("broj_dokumenta", anchor=W, width=120, stretch=True)
tree.column("br_ponude", anchor=W, width=150, stretch=True)
tree.column("Napomena", anchor=W, width=200, stretch=True)
tree.column("rowid", anchor=W, width=0, stretch=NO)

tree.heading("#0", text="\n", anchor=W)
tree.heading("br_naloga", text="BROJ", anchor=CENTER)
tree.heading("sektor", text="SEKTOR", anchor=CENTER)
tree.heading("izdavalac", text="IZDAVALAC", anchor=CENTER)
tree.heading("vođa_tima", text="VOĐA TIMA", anchor=CENTER)
tree.heading("član_tima_1", text="ČLAN 1", anchor=CENTER)
tree.heading("član_tima_2", text="ČLAN 2", anchor=CENTER)
tree.heading("član_tima_3", text="ČLAN 3", anchor=CENTER)
tree.heading("član_tima_4", text="ČLAN 4", anchor=CENTER)
tree.heading("klijent", text="KLIJENT", anchor=CENTER)
tree.heading("mesto", text="MESTO", anchor=CENTER)
tree.heading("adresa", text="ADRESA", anchor=CENTER)
tree.heading("aktivnost", text="AKTIVNOST", anchor=CENTER)
tree.heading("datum_izdavanja", text="IZDATO \nDANA", anchor=CENTER)
tree.heading("rok", text="ROK", anchor=CENTER)
tree.heading("kontakt", text="KONTAKT", anchor=CENTER)
tree.heading("broj_dokumenta", text="BROJ \nDOKUMENTA", anchor=CENTER)
tree.heading("br_ponude", text="BR \nPONUDE", anchor=CENTER)
tree.heading("Napomena", text="NAPOMENA", anchor=CENTER)
tree.heading("rowid", text="id", anchor=CENTER)

# bindings & events
imena = fetch_imena()
combo.bind('<<ComboboxSelected>>', sektor_changed)
populate_combobox(reciver_entry, imena)
populate_combobox(team_memeber_entry1, imena)
populate_combobox(team_memeber_entry2, imena)
populate_combobox(team_memeber_entry3, imena)
populate_combobox(team_memeber_entry4, imena)

reciver_entry.bind("<KeyRelease>", lambda event: filter_names(event, reciver_entry))
team_memeber_entry1.bind("<KeyRelease>", lambda event: filter_names(event, team_memeber_entry1))
team_memeber_entry2.bind("<KeyRelease>", lambda event: filter_names(event, team_memeber_entry2))
team_memeber_entry3.bind("<KeyRelease>", lambda event: filter_names(event, team_memeber_entry3))
team_memeber_entry4.bind("<KeyRelease>", lambda event: filter_names(event, team_memeber_entry4))

db_name = "Evidencija.db"

def get_available_years():
    conn = sq.connect("Evidencija.db")
    c = conn.cursor()
    c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'Evidencija%'")
    years = sorted([name.replace("Evidencija", "") for name, in c.fetchall() if name.replace("Evidencija", "").isdigit()], reverse=True)
    conn.close()
    return years

def get_table_name():
    return f"Evidencija{selected_year.get()}"

def update_treeview():
    """Update the treeview with data from the selected year's table."""
    table_name = get_table_name()
    conn = sq.connect(db_name)
    c = conn.cursor()
#potencijalno mogu ovde da mu dam komandu da sortira po broju naloga
    c.execute(f"SELECT *, rowid FROM {table_name} ORDER BY br_naloga")
    records = c.fetchall()

    # Clear the treeview before inserting new data
    tree.delete(*tree.get_children())

    # Add our data to the screen
    global count
    count = 0

    for record in records:
        if count % 2 == 0:
            tree.insert("", "end", iid=count, text='', values=record, tags=('evenrow',))

        else:
            tree.insert("", "end", iid=count, text='', values=record, tags=('oddrow',))

        # Increment counter inside the loop
        count += 1

    conn.commit()
    conn.close()

# Fetch years and set default year
available_years = get_available_years()
current_year = str(datetime.datetime.today().year)
selected_year = ttkb.StringVar(value=current_year if current_year in available_years else available_years[0])

# Year selection dropdown
ttkb.Label(basic_frame, text="Odaberi godinu:").grid(row=0, column=0, padx=5, pady=5)
year_selector = ttkb.Combobox(basic_frame, textvariable=selected_year, values=available_years, style="light")
year_selector.grid(row=0, column=1, padx=5, pady=5)
year_selector.bind("<<ComboboxSelected>>", lambda event: update_treeview())

# Load initial data
update_treeview()

for combo in combos:
    combo['values'] = imena

tree.bind("<ButtonRelease-1>", select_record)

client_entry.bind('<KeyRelease>', update_widgets)
client_entry.bind('<FocusOut>', fill_firma_on_focus_out)



#temporary here until determine why certain bug appears
print(f"maxnalog: {repr(maxnalog)}")

root.mainloop()
